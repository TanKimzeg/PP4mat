from typing import Any
from functools import singledispatch
import os
from docx import Document
from .logger import setup_logger
from . import utils
from collections import defaultdict
from config_converter.config_handle import Config
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph
import win32com.client

logger = setup_logger(__package__)
errors:dict[str,list[str]] = defaultdict(list)

def title_checker(document: DocumentObject, format_config: Config) -> None:
    logger.info("检查标题格式...")
    config = format_config.title_config
    if config is None:
        logger.warning("标题格式配置未提供，跳过检查。")
        return 
    raise NotImplementedError("标题格式检查功能尚未实现")
    # title_format = document.paragraphs[1]
    # print(document.paragraphs[0].text)
    # title_para_format = title_format.paragraph_format
    # title_font_format = document.paragraphs[1].runs[3].font
    # print(title_format.name)
    # print(title_font_format.name)
    # print(title_para_format.alignment.name)
    
    # 检查标题对齐格式
    # if title_para_format.alignment != config['alignment']:
    #     # logger.debug(f"Title alignment format: {title_para_format.alignment}")
    #     logger.error(f"Title alignment is incorrect: expected {config['alignment']}, found {title_para_format.alignment}")
    # else:
    #     logger.info("标题对齐格式正确")

    # # 检查标题字体大小
    # if utils.correct_size(document.paragraphs[0:1], 12):
    #     logger.info("标题字体大小正确")
    # else:
    #     logger.error("标题字体大小错误")

def heading_checker(document: DocumentObject, format_config: Config) -> None:
    global errors
    logger.info("检查小节标题格式...")
    raise NotImplementedError

def text_checker(document: DocumentObject, format_config: Config) -> None:
    global errors
    logger.info("检查正文格式...")
    config = format_config.text_config
    if config is None:
        logger.warning(f"正文格式配置未提供，跳过检查。")
        return
    # Placeholder for text checking logic
    raise NotImplementedError("正文格式检查功能尚未实现")

def formula_checker(document: DocumentObject, format_config: Config) -> None:
    global errors
    logger.info("检查公式格式...")
    config = format_config.formula_config
    if config is None:
        logger.warning(f"公式格式配置未提供，跳过检查。")
        return
    raise NotImplementedError("公式格式检查功能尚未实现")

@singledispatch # TODO:Bug to fix
def figure_checker(win32doc: Any) -> None:
    global errors
    logger.info("检查图片格式...")
    # config = format_config.figure_config
    # if config is None:
    #     logger.warning(f"图片格式配置未提供，跳过检查。")
    #     return
    # raise NotImplementedError("图片格式检查功能尚未实现")
    picture_cnt = 0
    for shape in win32doc.Shapes:
        # if shape.Type == win32com.client.constants.msoPicture:
        if shape.Type == 13:  # msoPicture
            picture_cnt += 1
            if shape.Anchor.Paragraphs.Count > 0:
                caption_p = shape.Anchor.Paragraphs(1)
                if not caption_p.Range.Text.strip().startswith(f"图{picture_cnt} "):
                    logger.error(f"图片{picture_cnt}的标题格式错误，应该以\"图{picture_cnt} \"开头，但实际为：{caption_p.Range.Text.strip()}")
                    errors["图片检测"].append(f"图片{picture_cnt}的标题格式错误，应该以\"图 {picture_cnt}\"开头")
                else:
                    logger.info(f"图片{picture_cnt}的标题格式正确：{caption_p.Range.Text.strip()}")
            else:
                logger.error(f"图片{picture_cnt}未找到对应的标题段落，请检查文档格式。")
                errors["图片检测"].append(f"图片{picture_cnt}未找到对应的标题段落")

@figure_checker.register(DocumentObject)
def _(document: DocumentObject) -> None:
    def is_img(run: Any) -> bool:
        drawing_elements = run._element.xpath('.//w:drawing')
        for drawing in drawing_elements:
            if drawing.xpath('.//a:graphic'):
                return True
        return False

    paragraphs = document.paragraphs
    picture_cnt = 0
    for i, p in enumerate(paragraphs):
        if p.runs and any(is_img(run) for run in p.runs):
            picture_cnt += 1
            if i + 1 < len(paragraphs):
                next_p = paragraphs[i + 1]
                if next_p.text.strip().startswith(f"图{picture_cnt} "):
                    logger.info(f"图片{picture_cnt}的标题格式正确：{next_p.text.strip()}")
                else:
                    logger.error(f"图片{picture_cnt}的标题格式错误，应该以\"图{picture_cnt} \"开头，但实际为：{next_p.text.strip()}")
                    errors["图片检测"].append(f"图片{picture_cnt}的标题格式错误，应该以\"图 {picture_cnt}\"开头")


def table_checker(document: DocumentObject, format_config: Config) -> None:
    global errors
    logger.info("检查表格格式...")
    # config = format_config.table_config
    # if config is None:
    #     logger.warning("表格格式配置未提供，跳过检查。")
    #     return
    # raise NotImplementedError("表格格式检查功能尚未实现")
    tables = document.tables
    paragraphs = document.paragraphs
    
    for i, table in enumerate(tables):
        table_index = None
        for j, p in enumerate(paragraphs):
            if p._element.getnext() == table._element:
                table_index = j
                break

        if table_index is not None and table_index > 0:
            header_p = paragraphs[table_index] 
            if not header_p.text.strip().startswith(f"表{i+1} "):
                logger.error(f"表格{i+1}的标题格式错误，应该以\"表{i+1} \"开头，但实际为：{header_p.text.strip()}")
                errors["表格检测"].append(f"表格{i+1}的标题格式错误，应该以\"表 {i+1}\"开头")
            else:
                logger.info(f"表格{i+1}的标题格式正确：{header_p.text.strip()}")
        else:
            logger.error(f"表格{i+1}未找到对应的标题段落，请检查文档格式。")
            errors["表格检测"].append(f"表格{i+1}未找到对应的标题段落")
    

def reference_checker(reference: list[Paragraph], format_config: Config) -> None:
    import re
    logger.info("检查参考文献格式...")
    config = format_config.reference_config
    if config is None:
        logger.warning("参考文献格式配置未提供，跳过检查。")
        return
    # Placeholder for reference checking logic
    pattern = re.compile(r'^\[\d+\]')
    cnt = 0
    en_cnt = 0
    for p in reference:
        if re.match(pattern, p.text.strip()):
            cnt += 1
            if any(char.isascii() for char in p.text):
                en_cnt += 1
    if cnt < config["min_count"]:
        logger.error(f"参考文献数量少于{config['min_count']}条，当前数量为{cnt}条，请检查！")
        errors["参考文献检测"].append(f"参考文献数量少于{config['min_count']}条，当前数量为{cnt}条")
    else:
        logger.info(f"参考文献数量满足要求，当前数量为{cnt}条。")
    
    # 检查英文参考文献数量
    if en_cnt < config["en_min_count"]:
        logger.error(f"英文参考文献数量少于{config['en_min_count']}条，当前数量为{en_cnt}条，请检查！")
        errors["参考文献检测"].append(f"英文参考文献数量少于{config['en_min_count']}条，当前数量为{en_cnt}条")
    else:
        logger.info(f"英文参考文献数量满足要求，当前数量为{en_cnt}条。")
    
def citation_count_checker(location: dict[str, list[Paragraph]], format_config: Config) -> None:
    logger.info("检查脚注数量")
    config = format_config.citation_min_count
    if config is None:
        logger.warning("脚注数量配置未提供，跳过检查。")
        return
    citation_count = utils.count_citations(location)
    if citation_count < config:
        logger.error(f"脚注数量少于{config}条，当前数量为{citation_count}条，请检查！")
        errors["脚注检测"].append(f"脚注数量少于{config}条，当前数量为{citation_count}条")

def section_checker(section_location: dict) -> None:
    global errors
    undergraduate_sections = [
        "毕业论文（设计）",
        "摘 要：", 
        "关键词：", 
        "Abstract:", 
        "Keywords:", 
        "目 录", 
        # "文献综述",
        "参考文献：", 
        "附  录："
    ]
    for section in undergraduate_sections:
        if section not in section_location or len(section_location[section]) == 0:
            logger.error(f"\"{section}\"缺失或位置不正确")
            errors["章节检测"].append(f"\"{section}\"缺失或位置不正确")
        else:
            logger.info(f"\"{section}\"部分存在 {len(section_location[section])} 个段落")

def survey_checker(document: DocumentObject, format_config: Config) -> None:
    global errors
    keywords = [
        "综述",
        "国内外",
        "现状"
    ]
    logger.info("检查文献综述部分...")
    found = False
    for p in document.paragraphs:
        if any(keyword in p.text for keyword in keywords):
            logger.info(f"找到文献综述相关内容：{p.text.strip()}")
            found = True
            return

    if not found:
        logger.error("文献综述部分缺失或不完整，请检查！")
        errors["文献综述检测"].append("文献综述部分缺失或不完整")



def page_checker(document: DocumentObject, format_config: Config) -> None:
    global errors
    logger.info("检查页面格式...")
    # Placeholder for page checking logic
    config = format_config.page_config
    if config is None:
        logger.warning(f"页面格式配置未提供，跳过检查。")
        return
    raise NotImplementedError("页面格式检查功能尚未实现")

def abstract_checker(abstact: list[Paragraph], format_config: Config) -> None:
    global errors
    logger.info("检查摘要格式...")
    config = format_config.abstract_config
    if config is None:
        logger.warning(f"摘要格式配置未提供，跳过检查。")
        return
    raise NotImplementedError("摘要格式检查功能尚未实现")

def check_enough_words(document: DocumentObject, format_config: Config) -> None:
    logger.info("检查文档字数...")
    config = format_config.words_min_count
    if config is None:
        logger.warning("字数配置未提供，跳过检查。")
        return
    word_count = utils.total_words(document)
    if word_count < config:
        logger.error(f"文档字数少于{config}字，当前字数为{word_count}字，请检查！")
        errors["字数检测"].append(f"文档字数少于{config}字，当前字数为{word_count}字")
    else:
        logger.info(f"文档字数满足要求，当前字数为{word_count}字。")

def check_cover_info(info: dict[str, str]) -> None:
    infomation = [
        "题目",
        "学号",
        "姓名",
        "学院",
        "年级",
        "专业",
        "区队",
        "指导教师",
    ]
    for info_key in infomation:
        if info_key not in info:
            logger.error(f"封面信息缺失：{info_key}，请检查！")
            errors["封面信息检测"].append(f"封面信息缺失：{info_key}")


def check_format(docx_path: str,format_config:Config) -> tuple[dict,dict]:
    global errors 
    
    # pywin32支持更复杂的Word文档操作,本项目中,论文的封面信息存储在文本框中,
    # 因此需要使用pywin32来提取文本框内容;检测图片和图题的关联也需要pywin32来实现.
    word = win32com.client.Dispatch("Word.Application")
    win32doc = word.Documents.Open(os.path.abspath(docx_path))
    cover_info = utils.cover_info_from_textbox(win32doc)
    figure_checker(win32doc)
    win32doc.Close()
    word.Quit()

    check_cover_info(cover_info)

    document = Document(docx_path)
    sections = utils.get_sections(document)
    section_checker(sections)
    survey_checker(document, format_config)
    table_checker(document, format_config)
    # figure_checker(document)
    reference_checker(sections["参考文献："], format_config)

    citation_count_checker(sections, format_config)
    check_enough_words(document, format_config)

    return errors, cover_info
